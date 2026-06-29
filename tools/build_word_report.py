"""Build the editable DOCX final report from the supplied SCUT template."""

from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(
    r"C:\Users\hp\xwechat_files\wxid_owx3mootag1922_1a09\msg\file\2026-06\Final Report Template.docx"
)
OUT_DIR = ROOT / "output" / "docx"
OUT_DOCX = OUT_DIR / "Steam_Review_Analysis_Final_Report_Formatted.docx"
RESULTS = ROOT / "output" / "results" / "local"
FIGURES = ROOT / "output" / "figures"
TMP = ROOT / "tmp"
LOGO = TMP / "scut_template_image.jpeg"

ACCENT = RGBColor(151, 36, 43)  # restrained SCUT-like red
DARK = RGBColor(31, 41, 55)
MUTED = RGBColor(90, 100, 115)
LIGHT_FILL = "F7F2F2"
HEADER_FILL = "97242B"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    total_dxa = int(sum(widths_in) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            row.cells[idx].width = Inches(width)

def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9DEE7")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_paragraph(doc: Document, text: str, style: str = "Normal", bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.add_run("- ").bold = True
        p.add_run(item)

def add_numbered(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run(f"{idx}. ").bold = True
        p.add_run(item)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    # Plain heading paragraph avoids the supplied template's list-formatted
    # built-in Heading styles, which produced visible square bullets.
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(13 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(15 if level == 1 else 12.5)
    run.font.color.rgb = ACCENT if level <= 2 else DARK
    run.bold = True

def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.45) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_small_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    set_table_widths(table, widths)
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], HEADER_FILL)
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=RGBColor(255, 255, 255))
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value)
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[i], LIGHT_FILL)
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    section = doc.sections[-1]
    section.top_margin = Inches(0.95)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT if name != "Heading 3" else DARK
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(5)


def extract_logo() -> None:
    TMP.mkdir(exist_ok=True)
    with ZipFile(TEMPLATE) as z:
        LOGO.write_bytes(z.read("word/media/image1.jpeg"))


def add_running_header(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(LOGO), width=Inches(1.2))
    label = p.add_run("   Steam Review Analysis | Big Data Final Project")
    label.font.size = Pt(8.5)
    label.font.color.rgb = MUTED

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    team = fp.add_run("Steam Insight Lab")
    team.font.size = Pt(8.5)
    team.font.color.rgb = MUTED

def fill_cover(doc: Document, metadata: dict) -> None:
    # Preserve the provided cover layout while filling the template fields.
    if doc.tables:
        table = doc.tables[0]
        values = {
            "Student Name": " / ".join(metadata["members"]),
            "Student ID": "",
            "Major": "Data Science and Big Data Technology",
            "Semester": "Second Semester of Academic Year 2025-2026",
        }
        for row in table.rows:
            key = row.cells[0].text.strip()
            for field, value in values.items():
                if field in key:
                    row.cells[1].text = value
        # The template uses underline-like fields. Center all visible content in
        # this information block so the cover looks balanced after filling.
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    extract_logo()

    metadata = json.loads((ROOT / "project_metadata.json").read_text(encoding="utf-8"))
    summary = json.loads((RESULTS / "summary.json").read_text(encoding="utf-8"))
    genre = pd.read_csv(RESULTS / "genre.csv")
    price = pd.read_csv(RESULTS / "price_band.csv")
    top_games = pd.read_csv(RESULTS / "top_reliable.csv")
    publishers = pd.read_csv(RESULTS / "publishers.csv")
    top_tags = pd.read_csv(RESULTS / "top_tags.csv")

    doc = Document(str(TEMPLATE))
    fill_cover(doc, metadata)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_styles(doc)
    add_running_header(section)

    add_heading(doc, "Steam Review Analysis", 1)
    add_paragraph(
        doc,
        "This report analyzes the Kaggle Steam Store Games dataset using a big data "
        "workflow designed around HDFS storage, PySpark preprocessing, Spark SQL "
        "aggregation, and Python-based visualization. The public dashboard is available at "
        "https://steam-review-analysis-lmeonleo.streamlit.app/.",
    )

    add_small_table(
        doc,
        ["Item", "Description"],
        [
            ["Dataset", "Steam Store Games by Nik Davis on Kaggle"],
            ["Records", f"{summary['clean_rows']:,} game-level records after cleaning"],
            ["Review outcomes", f"{summary['total_positive_ratings'] + summary['total_negative_ratings']:,} positive/negative rating outcomes"],
            ["Team", f"{metadata['team_name']} - {' / '.join(metadata['members'])}"],
            ["Tools", "HDFS, PySpark, Spark SQL, Pandas, Matplotlib/Seaborn, Streamlit"],
        ],
        [1.25, 4.75],
    )

    add_heading(doc, "1. Dataset Introduction", 1)
    add_paragraph(
        doc,
        "The selected topic is Steam Review Analysis. The source data comes from the "
        "real Kaggle dataset 'Steam Store Games', which contains game metadata, prices, "
        "genres, platforms, SteamSpy tags, and aggregate positive and negative review "
        "counts. Because the available review fields are aggregate ratings rather than "
        "individual text reviews, the project focuses on review outcome analysis: how "
        "many users responded positively or negatively, and how those outcomes vary by "
        "year, genre, price band, publisher, and tag.",
    )
    add_paragraph(
        doc,
        "The cleaned game table contains 27,075 games and 32.8 million review outcomes. "
        "The overall weighted positive rate is 82.58%, while the median number of review "
        "outcomes per game is 36, indicating a long-tail catalogue in which a small number "
        "of games receive very high attention.",
    )
    add_paragraph(
        doc,
        "Two source files are used most heavily in the analysis. The first is steam.csv, "
        "which acts as the main game-level fact table and contains release date, price, "
        "platform, genre, publisher, and positive/negative rating counts. The second is "
        "steamspy_tag_data.csv, which stores tag vote counts in a wide format and is "
        "normalized into a tag-level analytical table. Other raw files are retained as part "
        "of the dataset package, but the core questions can be answered from these two "
        "tables with a clear and reproducible pipeline.",
    )

    add_heading(doc, "2. HDFS Storage and Access Strategy", 1)
    add_paragraph(
        doc,
        "The project separates raw, curated, and analytical data zones so that the pipeline "
        "is reproducible and safe to rerun. Raw CSV files are copied into HDFS without "
        "modification, cleaned game records are written as partitioned Parquet, and final "
        "analysis outputs are materialized as compact CSV result tables for reporting and "
        "the dashboard.",
    )
    add_paragraph(
        doc,
        "This layout follows a common big-data lake pattern. The raw zone is append-only, "
        "which protects the original Kaggle files and makes it possible to audit every "
        "transformation. The curated zone stores typed and cleaned Parquet records, improving "
        "I/O efficiency compared with repeatedly scanning CSV. The result zone contains only "
        "small aggregates, so downstream visualization does not need to reprocess the full "
        "dataset.",
    )
    add_small_table(
        doc,
        ["Layer", "HDFS / Local Path", "Purpose"],
        [
            ["Raw zone", "/projects/steam/raw", "Immutable Kaggle CSV files"],
            ["Curated zone", "/projects/steam/curated/games", "Cleaned Parquet table partitioned by release year"],
            ["Result zone", "output/results/spark", "Aggregated CSV outputs from Spark SQL"],
            ["Web layer", "output/results/local", "Small pre-aggregated files read by Streamlit"],
        ],
        [1.05, 2.05, 2.9],
    )

    add_heading(doc, "3. Data Preprocessing", 1)
    add_paragraph(
        doc,
        "The preprocessing pipeline is implemented in PySpark with an explicit schema. "
        "Key steps include parsing release dates, converting price and rating fields to "
        "numeric types, handling missing or malformed values, deriving total review counts "
        "and positive-rate percentages, splitting semicolon-delimited genres, and exploding "
        "platform indicators for analysis.",
    )
    add_numbered(
        doc,
        [
            "Load CSV files from HDFS using explicit Spark schemas.",
            "Normalize release dates, prices, platforms, genres, publishers, and rating counts.",
            "Remove or repair invalid rows while preserving the original raw zone.",
            "Derive analytical fields: release year, review total, positive percentage, price band, and Wilson lower bound.",
            "Write curated records to Parquet and run Spark SQL aggregations.",
        ],
    )
    add_paragraph(
        doc,
        "A local Pandas-based verification script is also included. It is not used as a "
        "replacement for the PySpark pipeline; instead, it independently recomputes the "
        "reported aggregates on a machine without Hadoop. Matching results between the "
        "local verification output and the Spark design increases confidence that the "
        "business logic is correct.",
    )
    add_small_table(
        doc,
        ["Quality check", "Result", "Interpretation"],
        [
            ["Primary key", "27,075 unique games", "No duplicate game rows after cleaning"],
            ["Review fields", "32.8M outcomes", "Positive and negative counts are numeric and non-negative"],
            ["Release year", "1997-2019", "2019 is treated as an incomplete snapshot year"],
            ["Ranking eligibility", ">= 1,000 reviews", "Small samples are excluded from reliable top-game ranking"],
        ],
        [1.55, 1.55, 2.9],
    )

    add_heading(doc, "4. Analytical Questions and Methods", 1)
    add_bullets(
        doc,
        [
            "How did the Steam catalogue expand over time, and how did average reception change?",
            "Which genres combine broad coverage with reliable user reception?",
            "How do price bands relate to positive rates and review attention?",
            "Which games rank highest after accounting for uncertainty in small samples?",
            "Which publishers and tags dominate the reviewed catalogue?",
        ],
    )
    add_paragraph(
        doc,
        "Spark SQL is used for grouping, filtering, ranking, and producing final result "
        "tables. The ranking method uses the 95% Wilson lower confidence bound rather "
        "than raw positive rate alone, because a game with 3 positive reviews out of 3 "
        "should not outrank a game with hundreds of thousands of highly positive reviews.",
    )
    add_paragraph(
        doc,
        "The main analytical measures are total review outcomes, raw positive percentage, "
        "mean genre-level rating, median review attention, and Wilson lower-bound score. "
        "Together these measures separate popularity from reliability: review volume "
        "captures attention, while the Wilson score provides a conservative estimate of "
        "user satisfaction when sample sizes differ strongly across games.",
    )
    add_small_table(
        doc,
        ["Question", "Method", "Output"],
        [
            ["Catalogue growth", "Group by release_year", "yearly.csv and market evolution figure"],
            ["Genre reception", "Explode genres, aggregate rating and Wilson metrics", "genre.csv"],
            ["Price association", "Bucket numeric prices into bands", "price_band.csv"],
            ["Reliable ranking", "Filter by review volume and sort by Wilson bound", "top_reliable.csv"],
            ["Tag attention", "Unpivot SteamSpy tag columns", "top_tags.csv"],
        ],
        [1.6, 2.55, 1.85],
    )

    add_heading(doc, "5. Key Findings and Visualizations", 1)
    add_heading(doc, "5.1 Catalogue Growth", 2)
    add_paragraph(
        doc,
        "The number of games released per year rises sharply in the later years of the "
        "snapshot, especially after the mid-2010s. Mean positive reception remains within "
        "a relatively stable band, suggesting that the platform's growth reflects market "
        "expansion rather than a simple collapse in average user satisfaction.",
    )
    add_paragraph(
        doc,
        "This pattern is important for interpretation. If catalogue growth were accompanied "
        "by a dramatic fall in average ratings, it would suggest that easier publishing "
        "reduced observed quality. Instead, the result is more nuanced: the catalogue grows "
        "quickly, while rating averages move within a moderate range. The platform therefore "
        "appears to become broader and more diverse rather than simply worse.",
    )
    add_figure(doc, FIGURES / "01_market_evolution.png", "Figure 1. Steam catalogue growth and rating trend by release year.")

    add_heading(doc, "5.2 Genre-Level Reception", 2)
    add_paragraph(
        doc,
        "Genre analysis shows that genre labels overlap substantially, so one game can "
        "contribute to multiple categories. Wilson-adjusted genre scores make the comparison "
        "more conservative by considering both rating level and sample reliability.",
    )
    add_paragraph(
        doc,
        "Because many Steam games are tagged with multiple genres, genre-level results "
        "should be read as overlapping market segments rather than mutually exclusive "
        "categories. This is why the report emphasizes relative patterns and reliable "
        "reception instead of treating genres as isolated groups.",
    )
    add_figure(doc, FIGURES / "02_genre_quality.png", "Figure 2. Genre coverage and Wilson-adjusted reception.")

    add_small_table(
        doc,
        ["Genre", "Games", "Total reviews", "Mean rating %", "Mean Wilson %"],
        [
            [
                str(r["genre"]),
                f"{int(r['games']):,}",
                f"{int(r['total_reviews']):,}",
                f"{r['mean_rating_pct']:.2f}",
                f"{r['mean_wilson_pct']:.2f}",
            ]
            for _, r in genre.head(6).iterrows()
        ],
        [1.45, 0.75, 1.15, 1.25, 1.25],
    )

    add_heading(doc, "5.3 Price Bands", 2)
    add_paragraph(
        doc,
        "Price bands are interpreted descriptively rather than causally. Mid-priced games "
        "tend to receive more review attention than very low-priced games, while free games "
        "show broad catalogue coverage and high median review counts.",
    )
    add_paragraph(
        doc,
        "The price analysis does not prove that price causes higher or lower satisfaction. "
        "Different price bands may also differ by marketing budget, production quality, "
        "genre mix, and release strategy. However, the comparison is still useful because "
        "it shows how review attention and user reception are distributed across commercial "
        "segments of the Steam catalogue.",
    )
    add_figure(doc, FIGURES / "03_price_bands.png", "Figure 3. Rating level and review attention by price band.")
    add_small_table(
        doc,
        ["Price band", "Games", "Mean price", "Mean rating %", "Median reviews"],
        [
            [
                str(r["price_band"]),
                f"{int(r['games']):,}",
                f"{r['mean_price']:.2f}",
                f"{r['mean_rating_pct']:.2f}",
                f"{int(r['median_reviews']):,}",
            ]
            for _, r in price.iterrows()
        ],
        [1.05, 0.75, 0.95, 1.35, 1.35],
    )

    add_heading(doc, "5.4 Reliable Game Rankings", 2)
    best = top_games.iloc[0]
    add_paragraph(
        doc,
        f"The highest-ranked game by Wilson lower bound is {best['name']}. This confirms "
        "that the ranking is not simply rewarding small-sample perfect scores; it favours "
        "games with both strong reception and sufficient review volume.",
    )
    add_paragraph(
        doc,
        "This ranking is more suitable for recommendation-style interpretation than a raw "
        "positive-rate ranking. A plain percentage can be unstable when a game has very few "
        "reviews, while the Wilson lower bound penalizes uncertainty and rewards consistent "
        "performance at scale.",
    )
    add_figure(doc, FIGURES / "04_top_reliable.png", "Figure 4. Most reliable highly rated games by Wilson lower bound.")
    add_small_table(
        doc,
        ["Rank", "Game", "Reviews", "Rating %", "Wilson %"],
        [
            [
                str(i + 1),
                str(r["name"]),
                f"{int(r['total_reviews']):,}",
                f"{r['rating_pct']:.2f}",
                f"{r['wilson_pct']:.2f}",
            ]
            for i, (_, r) in enumerate(top_games.head(8).iterrows())
        ],
        [0.55, 2.65, 1.05, 0.9, 0.9],
    )

    add_heading(doc, "5.5 Publishers and Tags", 2)
    add_paragraph(
        doc,
        "Publisher portfolio analysis identifies publishers with both sufficient catalogue "
        "size and strong user reception. SteamSpy tags show the market's dominant content "
        "signals, with Action, Indie, Adventure, Multiplayer, and Singleplayer among the "
        "most-voted tags.",
    )
    add_small_table(
        doc,
        ["Publisher", "Games", "Total reviews", "Mean Wilson %"],
        [
            [
                str(r["publisher"]),
                f"{int(r['games']):,}",
                f"{int(r['total_reviews']):,}",
                f"{r['mean_wilson_pct']:.2f}",
            ]
            for _, r in publishers.head(6).iterrows()
        ],
        [2.1, 0.85, 1.45, 1.2],
    )
    add_small_table(
        doc,
        ["Tag", "Tagged games", "Total tag votes", "Mean rating %"],
        [
            [
                str(r["tag"]),
                f"{int(r['tagged_games']):,}",
                f"{int(r['total_tag_votes']):,}",
                f"{r['mean_rating_pct']:.2f}",
            ]
            for _, r in top_tags.head(10).iterrows()
        ],
        [1.45, 1.15, 1.45, 1.25],
    )

    add_heading(doc, "6. Dashboard Implementation", 1)
    add_paragraph(
        doc,
        "The optional dashboard is implemented with Streamlit and deployed to Streamlit "
        "Community Cloud. It reads only pre-aggregated result files, which keeps the web "
        "application lightweight while preserving the full big-data workflow behind the "
        "analysis. Public URL: https://steam-review-analysis-lmeonleo.streamlit.app/.",
    )
    add_paragraph(
        doc,
        "The dashboard is designed as a demonstration layer rather than the primary "
        "processing engine. It exposes the final aggregates through interactive tabs for "
        "market growth, genres, price bands, rankings, and methodology. This separation "
        "keeps the web application fast while making the HDFS/PySpark/Spark SQL workflow "
        "clear to evaluators.",
    )

    add_heading(doc, "7. Conclusion and Reflection", 1)
    add_paragraph(
        doc,
        "The project demonstrates a complete analytical workflow from raw data storage to "
        "cleaned distributed processing, SQL-based aggregation, visualization, interpretation, "
        "and web presentation. The strongest methodological choice is the use of Wilson "
        "lower-bound ranking, which makes the results more reliable than a raw percentage "
        "ranking. A limitation is that the dataset contains aggregate review outcomes rather "
        "than individual review texts, so sentiment analysis of written review content is "
        "outside the scope of this project.",
    )
    add_paragraph(
        doc,
        "Future work could extend the analysis in three directions. First, a larger review-text "
        "dataset could support NLP sentiment modeling and topic extraction. Second, joining "
        "Steam data with external sales or player-count data would help distinguish attention "
        "from commercial success. Third, a scheduled pipeline could refresh the dashboard "
        "automatically when new catalogue snapshots become available.",
    )

    add_heading(doc, "8. Team Roles", 1)
    add_small_table(
        doc,
        ["Member", "Main responsibilities"],
        [
            ["Gao Mingmin", "HDFS ingestion, PySpark ETL design, Spark SQL analysis, visualization, interpretation, report and dashboard"],
        ],
        [1.35, 4.65],
    )

    add_heading(doc, "References", 1)
    add_numbered(
        doc,
        [
            "Davis, N. Steam Store Games [Dataset]. Kaggle. https://www.kaggle.com/datasets/nikdavis/steam-store-games/data",
            "Apache Software Foundation. Apache Spark Documentation. https://spark.apache.org/docs/latest/",
            "Apache Software Foundation. Apache Hadoop HDFS Documentation. https://hadoop.apache.org/docs/",
            "Streamlit. Streamlit Documentation. https://docs.streamlit.io/",
        ],
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
